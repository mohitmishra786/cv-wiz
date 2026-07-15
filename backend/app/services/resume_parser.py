"""
Resume Parser Service
Extracts structured data from uploaded resume files (PDF/DOCX/TXT/MD)
Uses PyMuPDF (fitz) for PDF parsing and LLM for structuring

Enhanced with:
- Smart chunking for long resumes
- llama-3.3-70b-versatile model for better accuracy
- Comprehensive logging for debugging
- Retry logic with exponential backoff
- Shared AsyncGroq client for efficiency
"""

import io
import os
import re
import json
import time
import traceback
import asyncio
from typing import Dict, List, Optional, Any
from dataclasses import dataclass

from app.utils.logger import logger, log_function_call
from app.config import get_settings

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("[ResumeParser] PyMuPDF not installed, PDF parsing will be limited")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("[ResumeParser] python-docx not installed, DOCX parsing will be disabled")


# LLM Configuration
LLM_MODEL = "llama-3.3-70b-versatile"
# Larger chunks = fewer sequential/parallel Groq calls (model context is large).
MAX_CHUNK_CHARS = 6000
MAX_OUTPUT_TOKENS = 8000
MAX_RETRIES = 3
RETRY_BASE_DELAY = 1.0  # seconds
# Cap concurrent chunk extractions to avoid Groq rate limits while cutting wall time.
MAX_PARALLEL_CHUNK_EXTRACTS = 3

# Embedded document images (profile photos, logos in PDF/DOCX)
MAX_EXTRACTED_IMAGES = 5
MAX_IMAGE_BYTES = 400_000  # ~400KB raw bytes before base64
MIN_IMAGE_DIM = 48  # skip tiny icons
MAX_IMAGE_DIM = 2048


# Module-level shared AsyncGroq client (lazy initialization)
_groq_client = None


def get_groq_client() -> Optional[Any]:
    """Get or create the shared AsyncGroq client instance."""
    global _groq_client
    if _groq_client is None:
        settings = get_settings()
        if not settings.groq_api_key:
            return None
        from groq import AsyncGroq
        _groq_client = AsyncGroq(api_key=settings.groq_api_key)
        logger.info("[ResumeParser] Shared AsyncGroq client created")
    return _groq_client


@dataclass
class ExtractedExperience:
    """Structured experience data extracted from resume"""
    company: str
    title: str
    location: Optional[str]
    start_date: str
    end_date: Optional[str]
    current: bool
    description: str
    highlights: List[str]


@dataclass
class ExtractedEducation:
    """Structured education data extracted from resume"""
    institution: str
    degree: str
    field: str
    start_date: str
    end_date: Optional[str]
    gpa: Optional[float] = None


@dataclass
class ExtractedProject:
    """Structured project data extracted from resume"""
    name: str
    description: str
    technologies: List[str]
    url: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None


@dataclass
class ExtractedData:
    """Complete extracted resume data"""
    name: Optional[str]
    email: Optional[str]
    phone: Optional[str]
    summary: Optional[str]
    experiences: List[ExtractedExperience]
    education: List[ExtractedEducation]
    skills: List[str]
    projects: List[ExtractedProject]
    publications: Optional[List[Dict[str, Any]]] = None
    certifications: Optional[List[str]] = None


class ResumeParser:
    """
    Resume parsing service that extracts structured data from PDF/DOCX/TXT/MD files.
    Uses LLM with smart chunking for accurate extraction.
    """
    
    def __init__(self):
        self.settings = get_settings()
        logger.info("[ResumeParser] Initialized", {
            "pymupdf_available": PYMUPDF_AVAILABLE,
            "docx_available": DOCX_AVAILABLE,
            "groq_configured": bool(self.settings.groq_api_key),
            "llm_model": LLM_MODEL,
            "max_chunk_chars": MAX_CHUNK_CHARS,
        })
    
    @log_function_call
    async def parse_file(
        self, 
        file_input, 
        filename: str, 
        file_type: str = "resume",
        is_file_path: bool = False
    ) -> Dict[str, Any]:
        """
        Parse resume/cover letter file and extract structured data.
        
        Args:
            file_input: File content as bytes OR file path as string
            filename: Original filename for type detection
            file_type: "resume" or "cover-letter"
            is_file_path: If True, file_input is a file path; if False, it's bytes
            
        Returns:
            Dictionary containing extracted data
        """
        start_time = time.time()
        
        # If file_input is a path, read the file
        if is_file_path:
            if not isinstance(file_input, str):
                logger.error("[ResumeParser] is_file_path=True but file_input is not a string")
                return {"error": "Invalid file input: expected file path when is_file_path=True"}
            
            file_path = file_input
            file_size = os.path.getsize(file_path)
            
            logger.start_operation("resume_parse", {
                "filename": filename,
                "file_size": file_size,
                "file_type": file_type,
                "file_path": file_path,
            })
            
            # Read file content from path
            try:
                with open(file_path, 'rb') as f:
                    file_content = f.read()
            except Exception as e:
                logger.error("[ResumeParser] Failed to read file from path", {
                    "file_path": file_path,
                    "error": str(e),
                })
                return {"error": f"Failed to read file: {str(e)}"}
        else:
            # Legacy behavior: file_input is bytes
            if not isinstance(file_input, bytes):
                logger.error("[ResumeParser] is_file_path=False but file_input is not bytes")
                return {"error": "Invalid file input: expected bytes when is_file_path=False"}
            
            file_content = file_input
            
            logger.start_operation("resume_parse", {
                "filename": filename,
                "file_size": len(file_content),
                "file_type": file_type,
            })
        
        try:
            # Determine file type and extract text + embedded images
            filename_lower = filename.lower()
            images: List[Dict[str, Any]] = []

            if filename_lower.endswith('.pdf'):
                text = self._extract_text_from_pdf(file_content)
                images = self._extract_images_from_pdf(file_content)
            elif filename_lower.endswith(('.docx', '.doc')):
                text = self._extract_text_from_docx(file_content)
                images = self._extract_images_from_docx(file_content)
            elif filename_lower.endswith(('.txt', '.md', '.markdown')):
                text = self._extract_text_from_plaintext(file_content)
            else:
                logger.warning("[ResumeParser] Unsupported file type", {"filename": filename})
                return {"error": "Unsupported file type. Please upload PDF, DOCX, TXT, or MD."}
            
            logger.info("[ResumeParser] Text extracted successfully", {
                "text_length": len(text),
                "filename": filename,
                "images_found": len(images),
            })
            
            # Debug-only preview (not in production INFO logs)
            if text:
                logger.debug("[ResumeParser] Text preview", {"preview": text[:300]})
            
            if not text or len(text.strip()) < 50:
                logger.warning("[ResumeParser] Insufficient text extracted", {
                    "text_length": len(text) if text else 0,
                    "images_found": len(images),
                })
                # Image-only documents: still return images so UI can save profile photo
                if images:
                    return {
                        "error": (
                            "Could not extract enough text from the file. "
                            "Embedded images were extracted if present."
                        ),
                        "extracted_text_preview": text[:500] if text else None,
                        "images": images,
                        "extraction_method": "images-only",
                    }
                return {
                    "error": "Could not extract enough text from the file. The file may be empty, image-only, or corrupted.",
                    "extracted_text_preview": text[:500] if text else None,
                }
            
            # Cover letter: structured letter fields + same image payload as resume
            if file_type == "cover-letter":
                result = await self._parse_cover_letter(text)
                result["images"] = images
                if images:
                    result["profile_image"] = self._pick_profile_image(images)
                duration_ms = (time.time() - start_time) * 1000
                logger.end_operation("resume_parse", duration_ms, {
                    "filename": filename,
                    "file_type": "cover-letter",
                    "extraction_method": result.get("extraction_method", "unknown"),
                    "word_count": result.get("word_count"),
                    "images_found": len(images),
                })
                return result
            
            # Structure the extracted text using LLM
            structured_data = await self._structure_resume_text(text)
            structured_data["images"] = images
            if images:
                structured_data["profile_image"] = self._pick_profile_image(images)
            
            duration_ms = (time.time() - start_time) * 1000
            logger.end_operation("resume_parse", duration_ms, {
                "filename": filename,
                "extraction_method": structured_data.get("extraction_method", "unknown"),
                "experiences_found": len(structured_data.get("experiences", [])),
                "projects_found": len(structured_data.get("projects", [])),
                "skills_found": len(structured_data.get("skills", [])),
                "education_found": len(structured_data.get("education", [])),
                "images_found": len(images),
            })
            
            return structured_data
            
        except Exception as e:
            error_traceback = traceback.format_exc()
            logger.error("[ResumeParser] Parse failed with exception", {
                "filename": filename,
                "error_type": type(e).__name__,
                "error_message": str(e),
                "traceback": error_traceback,
            })
            logger.fail_operation("resume_parse", e, {"filename": filename})
            return {
                "error": f"Failed to parse resume: {str(e)}",
                "error_type": type(e).__name__,
                "traceback": error_traceback if self.settings.environment == "development" else None,
            }
    
    async def _parse_cover_letter(self, text: str) -> Dict[str, Any]:
        """Parse cover letter using LLM — same quality bar as resume extraction."""
        logger.info("[ResumeParser] Parsing cover letter with LLM")
        cleaned = text.strip()
        basic = {
            "content": cleaned,
            "word_count": len(cleaned.split()),
            "extraction_method": "basic",
            "name": None,
            "email": None,
            "phone": None,
            "job_title": None,
            "company_name": None,
            "recipient_name": None,
            "key_qualifications": [],
            "skills": [],
            "tone": "professional",
        }

        if not self.settings.groq_api_key:
            return basic

        try:
            client = get_groq_client()
            if not client:
                raise ValueError("Groq client not available")

            prompt = f"""Extract structured information from this cover letter. Return ONLY valid JSON:

{{
    "content": "The full cleaned cover letter body text (preserve paragraphs)",
    "name": "Applicant full name if present, or null",
    "email": "Applicant email if present, or null",
    "phone": "Applicant phone if present, or null",
    "recipient_name": "Hiring manager / recipient name if mentioned, or null",
    "company_name": "Target company name if mentioned, or null",
    "job_title": "Position being applied for if mentioned, or null",
    "key_qualifications": ["Key qualifications or achievements mentioned"],
    "skills": ["Skills explicitly mentioned"],
    "tone": "professional | casual | formal",
    "word_count": integer
}}

Cover Letter Text:
{text[:MAX_CHUNK_CHARS]}

Return ONLY JSON, no markdown or explanation."""

            response = await client.chat.completions.create(
                model=LLM_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a document parser. Extract structured data and "
                            "return valid JSON only. Never invent facts not in the text."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.1,
                max_tokens=3000,
            )

            if not response.choices or not response.choices[0].message.content:
                raise ValueError("Empty response from LLM")
            content = response.choices[0].message.content.strip()

            if content.startswith("```"):
                content = re.sub(r'^```\w*\n?', '', content)
                content = re.sub(r'\n?```$', '', content)

            result = json.loads(content)
            result["extraction_method"] = "llm"
            # Normalize aliases used by the frontend save path
            if result.get("sender_name") and not result.get("name"):
                result["name"] = result["sender_name"]
            if result.get("recipient_company") and not result.get("company_name"):
                result["company_name"] = result["recipient_company"]
            if not result.get("content"):
                result["content"] = cleaned
            if not result.get("word_count"):
                result["word_count"] = len(str(result.get("content", "")).split())

            logger.info("[ResumeParser] Cover letter parsed successfully", {
                "word_count": result.get("word_count"),
                "has_job_title": bool(result.get("job_title")),
                "has_company": bool(result.get("company_name")),
                "has_name": bool(result.get("name")),
            })

            return result

        except Exception as e:
            logger.warning("[ResumeParser] LLM cover letter parsing failed, using basic", {
                "error": str(e),
                "error_type": type(e).__name__,
            })
            return basic

    def _pick_profile_image(self, images: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
        """Prefer a square-ish photo (headshot) for profile image."""
        if not images:
            return None

        def score(img: Dict[str, Any]) -> float:
            w = int(img.get("width") or 0)
            h = int(img.get("height") or 0)
            if w <= 0 or h <= 0:
                return 0.0
            ratio = min(w, h) / max(w, h)
            # Prefer near-square, reasonably large photos
            size_score = min(w, h) / 400.0
            return ratio * 2.0 + min(size_score, 2.0)

        ranked = sorted(images, key=score, reverse=True)
        best = ranked[0]
        return {**best, "is_profile": True}

    def _image_record(
        self,
        raw: bytes,
        ext: str,
        width: int,
        height: int,
        *,
        page: Optional[int] = None,
    ) -> Optional[Dict[str, Any]]:
        """Build a base64 data-URL image record with size/dimension guards."""
        import base64

        if width < MIN_IMAGE_DIM or height < MIN_IMAGE_DIM:
            return None
        if width > MAX_IMAGE_DIM and height > MAX_IMAGE_DIM:
            # Likely a full-page scan — still keep if not huge on disk
            pass
        if len(raw) > MAX_IMAGE_BYTES:
            logger.debug("[ResumeParser] Skipping large image", {
                "bytes": len(raw),
                "width": width,
                "height": height,
            })
            return None

        ext_norm = (ext or "jpeg").lower().replace("jpg", "jpeg")
        mime = f"image/{ext_norm}" if ext_norm != "jpeg" else "image/jpeg"
        if ext_norm not in ("jpeg", "png", "gif", "webp", "bmp"):
            mime = "image/jpeg"
            ext_norm = "jpeg"

        b64 = base64.b64encode(raw).decode("ascii")
        return {
            "mime_type": mime,
            "ext": ext_norm,
            "width": width,
            "height": height,
            "byte_size": len(raw),
            "page": page,
            "is_profile": False,
            "data_url": f"data:{mime};base64,{b64}",
        }

    def _extract_images_from_pdf(self, file_content: bytes) -> List[Dict[str, Any]]:
        """Extract embedded images from a PDF via PyMuPDF."""
        if not PYMUPDF_AVAILABLE:
            return []

        images: List[Dict[str, Any]] = []
        seen_xrefs: set = set()
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            for page_index, page in enumerate(doc):
                for img in page.get_images(full=True):
                    xref = img[0]
                    if xref in seen_xrefs:
                        continue
                    seen_xrefs.add(xref)
                    try:
                        extracted = doc.extract_image(xref)
                    except Exception as extract_err:
                        # Skip corrupt/unsupported image streams; keep scanning others.
                        logger.debug(
                            "[ResumeParser] Skipping unreadable PDF image stream",
                            {
                                "xref": xref,
                                "error_type": type(extract_err).__name__,
                            },
                        )
                        continue
                    if not extracted or not extracted.get("image"):
                        continue
                    rec = self._image_record(
                        extracted["image"],
                        str(extracted.get("ext") or "jpeg"),
                        int(extracted.get("width") or 0),
                        int(extracted.get("height") or 0),
                        page=page_index + 1,
                    )
                    if rec:
                        images.append(rec)
                    if len(images) >= MAX_EXTRACTED_IMAGES:
                        break
                if len(images) >= MAX_EXTRACTED_IMAGES:
                    break
            doc.close()
            logger.info("[ResumeParser] PDF image extraction complete", {
                "images": len(images),
            })
        except Exception as e:
            logger.warning("[ResumeParser] PDF image extraction failed", {
                "error_type": type(e).__name__,
                "error": str(e),
            })
        return images

    def _extract_images_from_docx(self, file_content: bytes) -> List[Dict[str, Any]]:
        """Extract embedded images from a DOCX package."""
        if not DOCX_AVAILABLE:
            return []

        images: List[Dict[str, Any]] = []
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            doc = Document(io.BytesIO(file_content))
            for rel in doc.part.rels.values():
                if rel.reltype != RT.IMAGE:
                    continue
                try:
                    blob = rel.target_part.blob
                    content_type = getattr(rel.target_part, "content_type", "") or "image/jpeg"
                    ext = content_type.split("/")[-1].replace("jpeg", "jpeg").replace("jpg", "jpeg")
                    # Width/height unknown without decoding — use placeholders for scoring
                    width, height = 200, 200
                    if PYMUPDF_AVAILABLE:
                        try:
                            pix = fitz.Pixmap(blob)
                            width, height = pix.width, pix.height
                            if pix.n >= 5:  # CMYK etc.
                                pix = fitz.Pixmap(fitz.csRGB, pix)
                                blob = pix.tobytes("jpeg")
                                ext = "jpeg"
                            elif pix.alpha:
                                pix = fitz.Pixmap(fitz.csRGB, pix)
                                blob = pix.tobytes("png")
                                ext = "png"
                        except Exception as pix_err:
                            # Fall back to raw blob + placeholder dimensions.
                            logger.debug(
                                "[ResumeParser] DOCX image dimension probe failed",
                                {"error_type": type(pix_err).__name__},
                            )
                    rec = self._image_record(blob, ext, width, height)
                    if rec:
                        images.append(rec)
                    if len(images) >= MAX_EXTRACTED_IMAGES:
                        break
                except Exception as rel_err:
                    # Skip individual broken relationships; continue with others.
                    logger.debug(
                        "[ResumeParser] Skipping DOCX image relationship",
                        {"error_type": type(rel_err).__name__},
                    )
                    continue
            logger.info("[ResumeParser] DOCX image extraction complete", {
                "images": len(images),
            })
        except Exception as e:
            logger.warning("[ResumeParser] DOCX image extraction failed", {
                "error_type": type(e).__name__,
                "error": str(e),
            })
        return images
    
    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from a PDF file using PyMuPDF."""
        logger.debug("[ResumeParser] Extracting text from PDF")
        
        if not PYMUPDF_AVAILABLE:
            logger.error("[ResumeParser] PyMuPDF not available")
            raise ImportError("PyMuPDF is required for PDF parsing. Install with: pip install PyMuPDF")
        
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            text_parts = []
            
            for page_num, page in enumerate(doc):
                # Try multiple extraction methods for better results
                # Method 1: Standard text extraction
                page_text = page.get_text("text")
                
                # Method 2: If standard yields little text, try blocks
                if len(page_text.strip()) < 100:
                    blocks = page.get_text("blocks")
                    block_text = "\n".join([b[4] for b in blocks if isinstance(b[4], str)])
                    if len(block_text) > len(page_text):
                        page_text = block_text
                
                # Method 3: Try dict extraction for complex layouts
                if len(page_text.strip()) < 100:
                    dict_data = page.get_text("dict")
                    dict_text = ""
                    for block in dict_data.get("blocks", []):
                        if "lines" in block:
                            for line in block["lines"]:
                                for span in line.get("spans", []):
                                    dict_text += span.get("text", "") + " "
                                dict_text += "\n"
                    if len(dict_text) > len(page_text):
                        page_text = dict_text
                
                text_parts.append(page_text)
                logger.debug(f"[ResumeParser] Page {page_num + 1} extracted", {
                    "chars": len(page_text),
                })
            
            doc.close()
            full_text = "\n\n".join(text_parts)
            
            # Clean up PDF artifacts while preserving structure
            full_text = re.sub(r'[ \t]+', ' ', full_text)  # Normalize spaces
            full_text = re.sub(r'\n{3,}', '\n\n', full_text)  # Max 2 newlines
            full_text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', full_text)  # Fix hyphenation
            
            logger.info("[ResumeParser] PDF text extraction complete", {
                "total_pages": len(text_parts),
                "total_chars": len(full_text),
            })
            
            return full_text
            
        except Exception as e:
            error_traceback = traceback.format_exc()
            logger.error("[ResumeParser] PDF extraction failed", {
                "error": str(e),
                "error_type": type(e).__name__,
                "traceback": error_traceback,
            })
            raise
    
    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from a DOCX file."""
        logger.debug("[ResumeParser] Extracting text from DOCX")
        
        if not DOCX_AVAILABLE:
            logger.error("[ResumeParser] python-docx not available")
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also get text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts)
            
            logger.info("[ResumeParser] DOCX text extraction complete", {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "total_chars": len(full_text),
            })
            
            return full_text
            
        except Exception as e:
            error_traceback = traceback.format_exc()
            logger.error("[ResumeParser] DOCX extraction failed", {
                "error": str(e),
                "error_type": type(e).__name__,
                "traceback": error_traceback,
            })
            raise
    
    def _extract_text_from_plaintext(self, file_content: bytes) -> str:
        """Extract text from a plain text or markdown file."""
        logger.debug("[ResumeParser] Extracting text from plaintext/markdown")
        
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    logger.info("[ResumeParser] Plaintext extraction complete", {
                        "encoding": encoding,
                        "total_chars": len(text),
                    })
                    return text
                except UnicodeDecodeError:
                    continue
            
            text = file_content.decode('utf-8', errors='replace')
            logger.warning("[ResumeParser] Plaintext decoded with replacement characters")
            return text
            
        except Exception as e:
            error_traceback = traceback.format_exc()
            logger.error("[ResumeParser] Plaintext extraction failed", {
                "error": str(e),
                "error_type": type(e).__name__,
                "traceback": error_traceback,
            })
            raise
    
    def _chunk_resume_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Split resume into logical chunks for LLM processing.
        Each chunk contains complete sections to avoid cutting content mid-sentence.
        
        Returns list of dicts with 'text' and 'section_hints' keys.
        """
        logger.debug("[ResumeParser] Chunking resume text", {"total_chars": len(text)})
        
        # Section headers to look for
        section_patterns = [
            (r'(?i)\b(experience|work\s*history|employment|professional\s*experience)\b', 'experience'),
            (r'(?i)\b(education|academic|degrees?|qualifications?)\b', 'education'),
            (r'(?i)\b(skills?|technical\s*skills?|technologies?|proficiencies?|competencies?)\b', 'skills'),
            (r'(?i)\b(projects?|portfolio|personal\s*projects?|open\s*source)\b', 'projects'),
            (r'(?i)\b(publications?|papers?|research)\b', 'publications'),
            (r'(?i)\b(certifications?|certificates?|licenses?)\b', 'certifications'),
            (r'(?i)\b(summary|objective|profile|about\s*me)\b', 'summary'),
        ]
        
        # Find all section boundaries
        sections = []
        for pattern, section_type in section_patterns:
            for match in re.finditer(pattern, text):
                sections.append({
                    'type': section_type,
                    'start': match.start(),
                    'header': match.group()
                })
        
        # Sort sections by position
        sections.sort(key=lambda x: x['start'])
        
        # If text is short enough, return as single chunk
        if len(text) <= MAX_CHUNK_CHARS:
            logger.debug("[ResumeParser] Text fits in single chunk")
            return [{'text': text, 'section_hints': ['all']}]
        
        # If no sections found, split by character limit at sentence boundaries
        if not sections:
            logger.debug("[ResumeParser] No sections found, splitting by size")
            return self._split_by_size(text)
        
        # Create chunks based on sections
        chunks = []
        current_chunk = ""
        current_hints = []
        
        # Add header info (usually name, contact) as first chunk
        if sections and sections[0]['start'] > 100:
            header_text = text[:sections[0]['start']]
            if header_text.strip():
                chunks.append({
                    'text': header_text,
                    'section_hints': ['contact', 'summary']
                })
        
        # Process each section
        for i, section in enumerate(sections):
            # Get text until next section or end
            start = section['start']
            end = sections[i + 1]['start'] if i + 1 < len(sections) else len(text)
            section_text = text[start:end]
            
            # If adding this section exceeds limit, save current chunk and start new
            if len(current_chunk) + len(section_text) > MAX_CHUNK_CHARS:
                if current_chunk.strip():
                    chunks.append({
                        'text': current_chunk,
                        'section_hints': current_hints
                    })
                current_chunk = section_text
                current_hints = [section['type']]
            else:
                current_chunk += section_text
                current_hints.append(section['type'])
        
        # Don't forget the last chunk
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk,
                'section_hints': current_hints
            })
        
        logger.info("[ResumeParser] Resume chunked", {
            "total_chunks": len(chunks),
            "chunk_sizes": [len(c['text']) for c in chunks],
            "section_hints": [c['section_hints'] for c in chunks],
        })
        
        return chunks
    
    def _split_by_size(self, text: str) -> List[Dict[str, Any]]:
        """Split text by size at sentence boundaries."""
        chunks = []
        current = ""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        for sentence in sentences:
            if len(current) + len(sentence) > MAX_CHUNK_CHARS:
                if current.strip():
                    chunks.append({'text': current, 'section_hints': ['partial']})
                current = sentence
            else:
                current += " " + sentence if current else sentence
        
        if current.strip():
            chunks.append({'text': current, 'section_hints': ['partial']})
        
        return chunks
    
    async def _structure_resume_text(self, text: str) -> Dict[str, Any]:
        """
        Structure raw resume text into organized data using LLM.
        Uses chunking for long resumes and merges results.
        """
        logger.debug("[ResumeParser] Structuring resume text", {"text_length": len(text)})
        
        if not self.settings.groq_api_key:
            logger.warning("[ResumeParser] No Groq API key, falling back to regex")
            result = self._regex_extract(text)
            result["extraction_method"] = "regex"
            return result
        
        try:
            # Chunk the text
            chunks = self._chunk_resume_text(text)

            # Process chunks concurrently (bounded) — sequential 5-chunk parses
            # were taking ~2 minutes and timing out the Vercel proxy.
            semaphore = asyncio.Semaphore(MAX_PARALLEL_CHUNK_EXTRACTS)

            async def _process_chunk(index: int, chunk: Dict[str, Any]) -> Dict[str, Any]:
                async with semaphore:
                    logger.info(
                        f"[ResumeParser] Processing chunk {index + 1}/{len(chunks)}",
                        {
                            "chunk_size": len(chunk["text"]),
                            "section_hints": chunk["section_hints"],
                        },
                    )
                    return await self._llm_extract_chunk(
                        chunk["text"],
                        chunk["section_hints"],
                        index == 0,
                    )

            all_results = await asyncio.gather(
                *[_process_chunk(i, chunk) for i, chunk in enumerate(chunks)]
            )
            
            # Merge all results (order preserved by gather)
            merged = self._merge_chunk_results(list(all_results))
            merged["extraction_method"] = "llm"
            merged["chunks_processed"] = len(chunks)
            
            logger.info("[ResumeParser] Resume structuring complete", {
                "chunks_processed": len(chunks),
                "parallel_limit": MAX_PARALLEL_CHUNK_EXTRACTS,
                "experiences": len(merged.get("experiences", [])),
                "projects": len(merged.get("projects", [])),
                "skills": len(merged.get("skills", [])),
                "education": len(merged.get("education", [])),
            })
            
            return merged
            
        except Exception as e:
            error_traceback = traceback.format_exc()
            logger.warning("[ResumeParser] LLM extraction failed, falling back to regex", {
                "error": str(e),
                "error_type": type(e).__name__,
                "traceback": error_traceback,
            })
            result = self._regex_extract(text)
            result["extraction_method"] = "regex"
            return result
    
    async def _llm_extract_chunk(self, text: str, section_hints: List[str], is_first_chunk: bool) -> Dict[str, Any]:
        """Extract structured data from a single chunk using LLM with retry logic."""
        logger.start_operation("llm_chunk_extract")
        start_time = time.time()
        
        client = get_groq_client()
        if not client:
            raise ValueError("Groq client not available")
        
        # Build dynamic prompt based on section hints
        sections_to_extract = []
        if is_first_chunk or 'contact' in section_hints or 'summary' in section_hints or 'all' in section_hints:
            sections_to_extract.extend(['name', 'email', 'phone', 'summary'])
        if 'experience' in section_hints or 'all' in section_hints or 'partial' in section_hints:
            sections_to_extract.append('experiences')
        if 'education' in section_hints or 'all' in section_hints or 'partial' in section_hints:
            sections_to_extract.append('education')
        if 'skills' in section_hints or 'all' in section_hints or 'partial' in section_hints:
            sections_to_extract.append('skills')
        if 'projects' in section_hints or 'all' in section_hints or 'partial' in section_hints:
            sections_to_extract.append('projects')
        if 'publications' in section_hints or 'all' in section_hints:
            sections_to_extract.append('publications')
        if 'certifications' in section_hints or 'all' in section_hints:
            sections_to_extract.append('certifications')
        
        # If no specific hints, extract everything
        if not sections_to_extract:
            sections_to_extract = ['name', 'email', 'phone', 'summary', 'experiences', 'education', 'skills', 'projects']

        # sections_to_extract guides which fields we emphasize in the prompt
        sections_csv = ", ".join(sections_to_extract)
        
        prompt = f"""You are an expert resume parser. Extract ALL information from this resume section thoroughly.
Focus especially on these sections when present: {sections_csv}.
DO NOT miss any projects, experiences, skills, or contact information. Be exhaustive in your extraction.

CRITICAL EXTRACTION RULES:
1. NAME: Extract the candidate's FULL NAME exactly as written. Look at the first line or header of the resume.
2. EMAIL: Extract ANY email address found (format: user@domain.com)
3. PHONE: Extract ANY phone number found (including international formats)
4. PROJECT URLS: Extract GitHub links, portfolio links, live site URLs for each project. Look for github.com, gitlab.com, or any http/https URLs near project names.
5. PROJECT NAMES: Keep the EXACT project name as written on the resume, do not paraphrase.
6. PUBLICATIONS: Extract DOI links, arXiv links, or any URLs associated with publications.

Return ONLY a valid JSON object with this structure (include only relevant fields):

{{
    "name": "Full name EXACTLY as written on resume, or null",
    "email": "Complete email address, or null",
    "phone": "Complete phone number with country code if present, or null",
    "linkedin": "LinkedIn URL if present, or null",
    "github": "GitHub profile URL if present, or null",
    "portfolio": "Portfolio website URL if present, or null",
    "summary": "Professional summary or objective statement or null",
    "experiences": [
        {{
            "company": "Company name",
            "title": "Job title",
            "location": "Location or null",
            "start_date": "YYYY-MM or YYYY",
            "end_date": "YYYY-MM or YYYY or null if current",
            "current": true/false,
            "description": "Role description",
            "highlights": ["Achievement 1", "Achievement 2", ...]
        }}
    ],
    "education": [
        {{
            "institution": "School name",
            "degree": "Degree type",
            "field": "Field of study",
            "start_date": "YYYY",
            "end_date": "YYYY or null",
            "gpa": GPA as number or null
        }}
    ],
    "skills": ["Skill 1", "Skill 2", ...],
    "projects": [
        {{
            "name": "EXACT project name as written on resume",
            "description": "Detailed description",
            "technologies": ["Tech 1", "Tech 2", ...],
            "url": "GitHub URL, live site URL, or null",
            "github_url": "GitHub repository URL if available, or null",
            "live_url": "Live demo/site URL if available, or null",
            "start_date": "YYYY or null",
            "end_date": "YYYY or null"
        }}
    ],
    "publications": [
        {{
            "title": "Publication title EXACTLY as written",
            "venue": "Journal/Conference name",
            "year": "Year",
            "authors": ["Author 1", "Author 2"],
            "url": "DOI, arXiv, or publication URL, or null"
        }}
    ],
    "certifications": ["Certification 1", "Certification 2", ...],
    "achievements": ["Achievement 1", "Achievement 2", ...]
}}

Section hints: {', '.join(section_hints)}

IMPORTANT REMINDERS:
- Extract EVERY project mentioned with their URLs. Do not skip any.
- Extract the candidate's NAME, EMAIL, and PHONE from the header/contact section.
- Do not summarize project names - use EXACT names from the resume.
- Look for GitHub, LinkedIn, portfolio links and include them.

Resume Section:
{text}

Return ONLY the JSON object, no markdown formatting, no explanation."""

        last_error = None
        content = ""  # Initialize before try block
        for attempt in range(MAX_RETRIES):
            try:
                response = await client.chat.completions.create(
                    model=LLM_MODEL,
                    messages=[
                        {"role": "system", "content": "You are a precise resume parser. Extract all structured data exhaustively. Return valid JSON only. Never skip projects or experiences."},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.1,
                    max_tokens=MAX_OUTPUT_TOKENS,
                )
                
                content = response.choices[0].message.content.strip()
                
                # Clean up markdown formatting
                if content.startswith("```"):
                    content = re.sub(r'^```\w*\n?', '', content)
                    content = re.sub(r'\n?```$', '', content)
                
                result = json.loads(content)
                
                duration_ms = (time.time() - start_time) * 1000
                logger.end_operation("llm_chunk_extract", duration_ms, {
                    "model": LLM_MODEL,
                    "attempt": attempt + 1,
                    "tokens_used": response.usage.total_tokens if response.usage else None,
                    "experiences_found": len(result.get("experiences", [])),
                    "projects_found": len(result.get("projects", [])),
                })
                
                return result
                
            except json.JSONDecodeError as e:
                last_error = e
                logger.warning(f"[ResumeParser] JSON parse failed on attempt {attempt + 1}", {
                    "error": str(e),
                    "response_preview": content[:500] if content else None,
                })
                if attempt < MAX_RETRIES - 1:
                    await asyncio.sleep(RETRY_BASE_DELAY * (2 ** attempt))
                    
            except Exception as e:
                last_error = e
                logger.warning(f"[ResumeParser] LLM call failed on attempt {attempt + 1}", {
                    "error": str(e),
                    "error_type": type(e).__name__,
                })
                if attempt < MAX_RETRIES - 1:
                    await asyncio.sleep(RETRY_BASE_DELAY * (2 ** attempt))
        
        # All retries failed
        logger.error("[ResumeParser] All LLM extraction attempts failed", {
            "attempts": MAX_RETRIES,
            "last_error": str(last_error) if last_error else "No attempts performed",
        })
        if last_error:
            raise last_error
        else:
            raise RuntimeError("All LLM extraction attempts failed (no attempts performed)")
    
    def _merge_chunk_results(self, results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Merge results from multiple chunks into a single structured result."""
        merged = {
            "name": None,
            "email": None,
            "phone": None,
            "summary": None,
            "experiences": [],
            "education": [],
            "skills": [],
            "projects": [],
            "publications": [],
            "certifications": [],
        }
        
        for result in results:
            # Take first non-null values for single fields
            if not merged["name"] and result.get("name"):
                merged["name"] = result["name"]
            if not merged["email"] and result.get("email"):
                merged["email"] = result["email"]
            if not merged["phone"] and result.get("phone"):
                merged["phone"] = result["phone"]
            if not merged["summary"] and result.get("summary"):
                merged["summary"] = result["summary"]
            
            # Extend lists
            merged["experiences"].extend(result.get("experiences", []))
            merged["education"].extend(result.get("education", []))
            merged["projects"].extend(result.get("projects", []))
            merged["publications"].extend(result.get("publications", []))
            
            # Merge skills (deduplicate)
            for skill in result.get("skills", []):
                if skill and skill not in merged["skills"]:
                    merged["skills"].append(skill)
            
            # Merge certifications (deduplicate)
            for cert in result.get("certifications", []):
                if cert and cert not in merged["certifications"]:
                    merged["certifications"].append(cert)
        
        # Remove duplicates from lists based on key fields
        merged["experiences"] = self._deduplicate_list(merged["experiences"], "company", "title")
        merged["education"] = self._deduplicate_list(merged["education"], "institution", "degree")
        merged["projects"] = self._deduplicate_list(merged["projects"], "name")
        merged["publications"] = self._deduplicate_list(merged["publications"], "title")
        
        logger.debug("[ResumeParser] Chunks merged", {
            "experiences": len(merged["experiences"]),
            "projects": len(merged["projects"]),
            "skills": len(merged["skills"]),
            "education": len(merged["education"]),
        })
        
        return merged
    
    def _deduplicate_list(self, items: List[Dict], *key_fields: str) -> List[Dict]:
        """Remove duplicate items based on key fields."""
        seen = set()
        unique = []
        for item in items:
            key = tuple(str(item.get(f, "")).lower() for f in key_fields)
            if key not in seen:
                seen.add(key)
                unique.append(item)
        return unique
    
    def _regex_extract(self, text: str) -> Dict[str, Any]:
        """Fallback regex-based extraction for basic resume parsing."""
        logger.info("[ResumeParser] Using regex extraction (LLM unavailable or failed)")
        
        result = {
            "name": None,
            "email": None,
            "phone": None,
            "summary": None,
            "experiences": [],
            "education": [],
            "skills": [],
            "projects": [],
            "publications": [],
            "certifications": [],
            "warning": "Using basic regex extraction. Results may be incomplete. For better results, ensure GROQ_API_KEY is configured.",
        }
        
        # Extract name (first line or capitalized words)
        lines = text.strip().split('\n')
        if lines:
            first_line = lines[0].strip()
            if re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+){1,3}$', first_line):
                result["name"] = first_line
                logger.debug("[ResumeParser] Name found")
        
        # Extract email
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        if email_match:
            result["email"] = email_match.group()
            logger.debug("[ResumeParser] Email found")
        
        # Extract phone
        phone_patterns = [
            r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',
            r'\+?[0-9]{1,3}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}',
        ]
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                result["phone"] = phone_match.group().strip()
                logger.debug("[ResumeParser] Phone found")
                break
        
        # Extract skills
        skills_patterns = [
            r'(?:skills?|technical skills?|technologies?|proficiencies?)[:\s]*([^\n]+(?:\n(?![A-Z][a-z]+:)[^\n]+)*)',
            r'(?:programming|languages?)[:\s]*([^\n]+)',
        ]
        for pattern in skills_patterns:
            skills_section = re.search(pattern, text, re.IGNORECASE)
            if skills_section:
                skills_text = skills_section.group(1)
                skills = re.split(r'[,|•·;\n]', skills_text)
                skills = [s.strip() for s in skills if s.strip() and len(s.strip()) < 50 and len(s.strip()) > 1]
                result["skills"].extend(skills[:30])
                break
        
        result["skills"] = list(dict.fromkeys(result["skills"]))
        
        logger.info("[ResumeParser] Regex extraction complete", {
            "has_name": bool(result["name"]),
            "has_email": bool(result["email"]),
            "has_phone": bool(result["phone"]),
            "skills_count": len(result["skills"]),
        })
        
        return result


# Singleton instance
resume_parser = ResumeParser()
